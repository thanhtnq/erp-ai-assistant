"""
Debug image mapping for ingest_knowledge.py
Run: python debug_images.py
"""
from pathlib import Path
from docx import Document
import re

FILE_PATH = "documents/_global/Sales/GLOBE3 ERP MANUAL 2019 - SALES VERSION 2.1.docx"
NS_REL    = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

doc  = Document(FILE_PATH)
body = doc.element.body

# ── 1. Build rId → image name map ─────────────────────────────────────────────
rid_to_image = {}
for rel in doc.part.rels.values():
    if "image" in rel.reltype.lower():
        rid_to_image[rel.rId] = Path(rel.target_ref).name

print(f"[1] Total rId→image mappings: {len(rid_to_image)}")

# ── 2. Find body elements containing images ────────────────────────────────────
image_map = {}  # body_idx → [image_name]
for idx, el in enumerate(body):
    tag  = el.tag.split("}")[-1]
    rids = []
    for child in el.iter():
        for attr_val in child.attrib.values():
            if attr_val.startswith("rId") and attr_val in rid_to_image:
                if attr_val not in rids:
                    rids.append(attr_val)
    if rids:
        image_map[idx] = [rid_to_image[r] for r in rids]

print(f"[2] Body elements with images: {len(image_map)}")
print("    First 5:")
for i, (bi, imgs) in enumerate(sorted(image_map.items())[:5]):
    print(f"    body[{bi}] → {imgs}")

# ── 3. Build paragraph → body index map ───────────────────────────────────────
para_body_idx = {}
for body_idx, child in enumerate(body):
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "p":
        para_body_idx[id(child)] = body_idx

print(f"\n[3] Paragraphs mapped to body index: {len(para_body_idx)}")

# ── 4. Find headings and their body indexes ────────────────────────────────────
RE = re.compile(r"^(\d+(?:\.\d+)*)[\s\t]+(.+)$")
print("\n[4] Heading paragraphs and their body indexes:")
for para in doc.paragraphs:
    text       = para.text.replace("\t", " ").strip()
    style_name = (para.style.name or "").lower()
    m          = RE.match(text)
    is_bold    = any(r.bold for r in para.runs if r.text.strip())

    if m and ("heading" in style_name or is_bold):
        bi = para_body_idx.get(id(para._p), -1)
        print(f"    body[{bi:4}] [{m.group(1):6}] style={para.style.name!r:15} {m.group(2)[:50]}")

# ── 5. Check if image body indexes overlap with section body indexes ───────────
print("\n[5] Sample body index range check:")
print(f"    Image body indexes (first 10): {sorted(image_map.keys())[:10]}")
heading_idxs = []
for para in doc.paragraphs:
    text       = para.text.replace("\t", " ").strip()
    style_name = (para.style.name or "").lower()
    m          = RE.match(text)
    is_bold    = any(r.bold for r in para.runs if r.text.strip())
    if m and ("heading" in style_name or is_bold):
        bi = para_body_idx.get(id(para._p), -1)
        if bi >= 0:
            heading_idxs.append(bi)

print(f"    Heading body indexes (first 10): {heading_idxs[:10]}")

# ── 6. Try nearest-section logic manually ─────────────────────────────────────
print("\n[6] Manual nearest-section test:")
# Simulate body_to_section_idx from content paragraphs
body_to_sec = {}
sec_idx = 0
for para in doc.paragraphs:
    text       = para.text.replace("\t", " ").strip()
    style_name = (para.style.name or "").lower()
    m          = RE.match(text)
    is_bold    = any(r.bold for r in para.runs if r.text.strip())
    if m and ("heading" in style_name or is_bold):
        sec_idx += 1
    bi = para_body_idx.get(id(para._p), -1)
    if bi >= 0 and text:
        body_to_sec[bi] = sec_idx

sorted_keys = sorted(body_to_sec.keys())

for img_bi in sorted(image_map.keys())[:5]:
    best = None
    for k in sorted_keys:
        if k <= img_bi:
            best = body_to_sec[k]
        else:
            break
    print(f"    image body[{img_bi}] → section {best} | images={image_map[img_bi]}")